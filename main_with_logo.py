
import streamlit as st
import pandas as pd
import networkx as nx
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from PIL import Image

# Tampilkan logo MSBC di atas semua halaman
logo = Image.open("logo_msbc.png")
st.image(logo, width=120)
st.markdown("## Aplikasi Sosiometri MSBC – PAUD sampai SMA")
st.markdown("_Disusun oleh: Mini Survival Boarding Camp (MSBC)_")

def login_page():
    st.markdown("### 🔐 Login Pengguna")
    username = st.text_input("Nama Pengguna")
    password = st.text_input("Kata Sandi", type="password")
    role = st.selectbox("Pilih Jenjang yang Diampu", ["PAUD", "SD", "SMP", "SMA"])

    if st.button("Masuk"):
        if username and password:
            st.success(f"Selamat datang, {username}! Anda login sebagai guru jenjang {role}.")
            st.session_state["logged_in"] = True
            st.session_state["role"] = role
        else:
            st.warning("Silakan isi nama pengguna dan kata sandi.")

if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if not st.session_state["logged_in"]:
    login_page()
else:
    role = st.session_state["role"]
    st.title(f"📊 Aplikasi Sosiometri Jenjang {role}")

    pertanyaan_dict = {
        "PAUD": [
            "Siapa teman yang paling enak diajak ngobrol atau bercerita?",
            "Siapa teman yang suka bantu kamu kalau kamu sedih atau takut?",
            "Kalau main bareng, siapa teman yang bisa tenang dan nggak mudah marah?",
            "Siapa teman yang paling kamu percaya buat jadi pemimpin permainan atau kegiatan?",
            "Kalau kita jalan-jalan atau main di luar, kamu paling suka bareng siapa?",
            "Siapa teman yang sering salat, berdoa, atau bantu beresin mainan menurut kamu?"
        ],
        "SD": [
            "Siapa teman yang sering bermain bersama kamu saat istirahat?",
            "Siapa teman yang kamu percaya saat kerja kelompok?",
            "Siapa teman yang selalu menghiburmu saat sedih?",
            "Siapa teman yang kamu sukai karena jujur dan adil?",
            "Siapa teman yang aktif dan suka membantu guru?",
            "Siapa teman yang bisa menenangkan saat marah?"
        ],
        "SMP": [
            "Siapa teman yang bisa kamu ajak diskusi serius?",
            "Siapa teman yang menurutmu memiliki jiwa kepemimpinan?",
            "Siapa teman yang tidak menyebarkan gosip dan bisa dipercaya?",
            "Siapa teman yang menunjukkan empati saat kamu dalam masalah?",
            "Siapa teman yang paling kamu hormati karena karakternya?",
            "Siapa teman yang berani membela temannya saat dibully?"
        ],
        "SMA": [
            "Siapa teman yang kamu anggap paling suportif?",
            "Siapa teman yang bisa kamu jadikan partner belajar?",
            "Siapa teman yang paling aktif di kegiatan sosial sekolah?",
            "Siapa teman yang bisa dipercaya menjaga rahasia?",
            "Siapa teman yang mampu memimpin dengan bijaksana?",
            "Siapa teman yang menurutmu punya integritas tinggi?"
        ]
    }

    st.markdown("### 📄 Daftar Pertanyaan Sosiometri")
    for i, q in enumerate(pertanyaan_dict[role], 1):
        st.markdown(f"**{i}. {q}**")

    uploaded_file = st.file_uploader("📥 Upload jawaban sosiometri dalam format Excel (Pilihan 1 - 18)", type=["xlsx"])

    def interpretasi_skor(skor):
        if skor == 0: return "❗ Terisolasi – Perlu perhatian khusus"
        elif skor <= 2: return "⚠️ Sosial Terbatas – Perlu dorongan interaksi"
        elif skor <= 5: return "✅ Cukup Sosial"
        elif skor <= 9: return "🌟 Populer – Disukai banyak teman"
        else: return "🏅 Sangat Populer – Potensial jadi fasilitator"

    def proses_sosiometri(df):
        nama_siswa = df['Nama Siswa'].tolist()
        popularitas = {nama: 0 for nama in nama_siswa}
        hubungan = []
        pilihan_cols = [col for col in df.columns if col.startswith("Pilihan")]
        for _, row in df.iterrows():
            pemilih = row['Nama Siswa']
            for kolom in pilihan_cols:
                dipilih = row[kolom]
                if pd.notna(dipilih) and dipilih in popularitas:
                    popularitas[dipilih] += 1
                    hubungan.append((pemilih, dipilih))
        df_hasil = pd.DataFrame([
            {"Nama": nama, "Skor Popularitas": skor, "Interpretasi": interpretasi_skor(skor)}
            for nama, skor in popularitas.items()
        ])
        df_hasil = df_hasil.sort_values(by="Skor Popularitas", ascending=False).reset_index(drop=True)
        return df_hasil, hubungan

    def simpan_word(df):
        doc = Document()
        doc.add_heading("Hasil Sosiometri dan Interpretasi", 0)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nama'
        hdr_cells[1].text = 'Skor Popularitas'
        hdr_cells[2].text = 'Interpretasi'
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Nama'])
            row_cells[1].text = str(row['Skor Popularitas'])
            row_cells[2].text = str(row['Interpretasi'])
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def simpan_gambar_sosiogram(hubungan, df_hasil):
        G = nx.DiGraph()
        G.add_edges_from(hubungan)
        node_colors = []
        labels = {}
        for i, row in df_hasil.iterrows():
            nama = row['Nama']
            skor = row['Skor Popularitas']
            label = f"{i+1}. {nama}"
            labels[nama] = label
            if skor == 0:
                node_colors.append("red")
            elif skor <= 2:
                node_colors.append("orange")
            elif skor <= 5:
                node_colors.append("yellow")
            else:
                node_colors.append("green")
        pos = nx.spring_layout(G, seed=42)
        fig, ax = plt.subplots(figsize=(10, 8))
        nx.draw(G, pos, with_labels=False, arrows=True, node_color=node_colors, node_size=2500, edge_color='gray')
        nx.draw_networkx_labels(G, pos, labels=labels, font_size=10)
        buf = BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        return buf, fig

    if uploaded_file:
        df_excel = pd.read_excel(uploaded_file)
        try:
            df_hasil, hubungan = proses_sosiometri(df_excel)
            st.success("✅ Data berhasil diproses!")
            st.dataframe(df_hasil)

            st.download_button("⬇️ Download Word", data=simpan_word(df_hasil), file_name="hasil_sosiometri.docx")

            png_file, fig = simpan_gambar_sosiogram(hubungan, df_hasil)
            st.pyplot(fig)
            st.download_button("⬇️ Download Gambar PNG", data=png_file, file_name="sosiogram.png", mime="image/png")
        except Exception as e:
            st.error(f"Terjadi kesalahan: {e}")
